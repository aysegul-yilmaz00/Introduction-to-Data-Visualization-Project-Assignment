from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import config
import os
import json
import re
from datetime import datetime

def parse_cevap(cevap_metni):
    """Gelen metni JSON olarak ayrıştırır."""
    match = re.search(r'```json\s*(.*?)\s*```', cevap_metni, re.DOTALL)
    if match:
        json_str = match.group(1)
    else:
        json_str = cevap_metni
        
    try:
        data = json.loads(json_str)
        return data
    except Exception as e:
        return {
            "makam": "İLGİLİ MAKAMA",
            "konu": "Konu Belirtilmedi",
            "ilgi": "",
            "icerik": cevap_metni
        }

def word_kaydet(metin, dosya_adi):
    """Verilen metni/JSON'u resmi yazı (şablon) formatında bir Word belgesine kaydeder."""
    data = parse_cevap(metin)

    doc = Document()
    
    # 1. Başlık Alanı (Kurum Adı)
    baslik = doc.add_paragraph()
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = baslik.add_run(config.KURUM_ADI)
    run.bold = True
    run.font.size = Pt(12)
    
    # Birim Adı
    birim = doc.add_paragraph()
    birim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_birim = birim.add_run(config.BIRIM_ADI)
    run_birim.font.size = Pt(11)

    doc.add_paragraph("\n") # Boşluk
    
    # 2. Sayı, Konu ve Tarih
    table = doc.add_table(rows=2, cols=3)
    table.autofit = True
    
    # Satır 0: Sayı ve Tarih
    row0 = table.rows[0].cells
    row0[0].text = "Sayı :"
    
    import random
    zaman_etiketi = datetime.now().strftime("%Y%m%d")
    rastgele_sayi = random.randint(1000, 9999)
    otomatik_sayi = f"E-34567890-010.99-{zaman_etiketi}{rastgele_sayi}"
    row0[1].text = otomatik_sayi
    
    bugun = datetime.now().strftime("%d.%m.%Y")
    p0_right = row0[2].paragraphs[0]
    p0_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p0_right.add_run(bugun).font.size = Pt(12)
    
    # Satır 1: Konu
    row1 = table.rows[1].cells
    row1[0].text = "Konu :"
    row1[1].text = data.get('konu', '')

    doc.add_paragraph("\n")
    
    # 3. Muhatap Makam
    makam = doc.add_paragraph()
    makam.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_makam = makam.add_run(data.get('makam', 'İLGİLİ MAKAMA'))
    run_makam.bold = True
    run_makam.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 4. İlgi (Varsa)
    ilgi_metni = data.get('ilgi', '').strip()
    if ilgi_metni:
        ilgi = doc.add_paragraph()
        r1 = ilgi.add_run("İlgi : ")
        r1.bold = True
        ilgi.add_run(ilgi_metni)
        doc.add_paragraph()

    # 5. İçerik (Metin)
    icerik_metni = data.get('icerik', '')
    paragraf = doc.add_paragraph()
    paragraf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraf.paragraph_format.first_line_indent = Cm(1.25)
    run_icerik = paragraf.add_run(icerik_metni)
    run_icerik.font.name = 'Times New Roman'
    run_icerik.font.size = Pt(12)
    
    doc.add_paragraph("\n\n")
    
    # 6. İmza Alanı (Sağa dayalı)
    imza_paragrafi = doc.add_paragraph()
    imza_paragrafi.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    imza_paragrafi.add_run("[İmza Sahibi Adı Soyadı]\n").bold = True
    imza_paragrafi.add_run("Birim Amiri / Müdür")

    # Dosyayı Kaydet
    kayit_yolu = os.path.join(config.CEVAPLAR_DIR, dosya_adi)
    doc.save(kayit_yolu)
    return kayit_yolu
